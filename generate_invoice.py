from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_receipt_docx(path: str) -> None:
    today_str = "2026-03-01"

    document = Document()

    # Header: Business name and phone
    title = document.add_paragraph()
    run = title.add_run("Fidinsky Tech Solutions")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = document.add_paragraph()
    p.add_run("Phone: +263779964877")

    document.add_paragraph()  # spacer

    # Receipt label
    p = document.add_paragraph()
    run = p.add_run("RECEIPT")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Receipt details
    details = document.add_paragraph()
    details.add_run(f"Date: {today_str}\n").bold = True
    details.add_run("Status: PAID (Cash)").bold = True

    document.add_paragraph()  # spacer

    # Client section
    bill_to_title = document.add_paragraph()
    bill_to_title.add_run("Client (Bill To)").bold = True

    bill_to = document.add_paragraph()
    bill_to.add_run("Client Name: BoltLog")

    document.add_paragraph()  # spacer

    # Line items table
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Qty"
    hdr_cells[3].text = "Rate"
    hdr_cells[4].text = "Amount"

    # Row 1: Android development (paid in full for agreed scope)
    row_cells = table.add_row().cells
    row_cells[0].text = "1"
    row_cells[1].text = "Android application development (paid in full)"
    row_cells[2].text = "1"
    row_cells[3].text = "375"
    row_cells[4].text = "375"

    # Row 2: Services integration
    row_cells = table.add_row().cells
    row_cells[0].text = "2"
    row_cells[1].text = "Services integration"
    row_cells[2].text = "1"
    row_cells[3].text = "150"
    row_cells[4].text = "150"

    document.add_paragraph()  # spacer

    # Total
    totals = document.add_paragraph()
    totals.add_run("Total Paid: 525").bold = True

    document.add_paragraph()  # spacer

    # Payment details
    pay_title = document.add_paragraph()
    pay_title.add_run("Payment Details").bold = True

    pay = document.add_paragraph()
    pay.add_run("Payment Method: Cash\n")
    pay.add_run("Amount Received: 525\n")
    pay.add_run(f"Date Received: {today_str}")

    document.add_paragraph()  # spacer

    # Notes
    notes_title = document.add_paragraph()
    notes_title.add_run("Notes").bold = True

    notes = document.add_paragraph()
    notes.add_run(
        "Project / Reference: Android application development and services integration (totals $525).\n"
    )
    notes.add_run("Additional Notes: Thank you for your business.")

    document.save(path)


if __name__ == "__main__":
    output_path = "boltlog_receipt_2026-03-01.docx"
    create_receipt_docx(output_path)
